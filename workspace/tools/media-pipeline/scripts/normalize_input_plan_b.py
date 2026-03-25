#!/home/user/.openclaw/workspace/tools/media-pipeline/.venv/bin/python
import argparse
import json
import re
import unicodedata
import urllib.request
from html import unescape
from pathlib import Path


CONTROL_CHAR_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]')
HTML_TAG_RE = re.compile(r'(?is)<[^>]+>')
FURIGANA_BRACE_RE = re.compile(r'\{([^{}:/]+)[/:][^{}]+\}')


def split_long_block(block: str, max_chars: int = 220):
    block = block.strip()
    if not block or len(block) <= max_chars:
        return [block] if block else []

    sentences = re.split(r'(?<=[。！？!?…」』）\)])\s*', block)
    sentences = [s.strip() for s in sentences if s.strip()]
    if len(sentences) <= 1:
        return [block[i:i + max_chars].strip() for i in range(0, len(block), max_chars) if block[i:i + max_chars].strip()]

    out = []
    current = ''
    for sent in sentences:
        candidate = sent if not current else current + sent
        if current and len(candidate) > max_chars:
            out.append(current.strip())
            current = sent
        else:
            current = candidate
    if current.strip():
        out.append(current.strip())
    return out


def strip_furigana_braces(text: str) -> str:
    prev = None
    while prev != text:
        prev = text
        text = FURIGANA_BRACE_RE.sub(r'\1', text)
    text = re.sub(r'\{[^{}]*\}', '', text)
    return text


def collapse_duplicate_lines(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    out = []
    for line in lines:
        if not line:
            if out and out[-1] != '':
                out.append('')
            continue
        if out and line == out[-1]:
            continue
        out.append(line)
    return '\n'.join(out)


def dedupe_repeated_prefix_progression(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if len(lines) < 3:
        return text
    if all(lines[i] in lines[i + 1] for i in range(len(lines) - 1)):
        return lines[-1]
    if all(lines[i + 1] in lines[i] for i in range(len(lines) - 1)):
        return lines[0]
    return text


def dedupe_repeated_sentences(text: str) -> str:
    parts = re.split(r'(?<=[。！？!?…])', text)
    parts = [p for p in parts if p]
    if len(parts) < 2:
        return text
    out = []
    prev_norm = None
    for part in parts:
        norm = re.sub(r'\s+', '', part)
        if norm and norm == prev_norm:
            continue
        out.append(part)
        prev_norm = norm or prev_norm
    return ''.join(out)


def heavy_cleanup_text(text: str) -> str:
    if not text:
        return ''
    text = unicodedata.normalize('NFKC', text)
    text = text.replace('\r\n', '\n').replace('\r', '\n').replace('\u00a0', ' ')
    text = CONTROL_CHAR_RE.sub('', text)
    text = re.sub(r'[\t\f\v]+', ' ', text)
    text = strip_furigana_braces(text)
    text = HTML_TAG_RE.sub(' ', text)
    text = unescape(text)
    text = re.sub(r'[♥❤♡🖤💗💓💞💕💘]+', '', text)
    text = re.sub(r'[※★☆◆◇□■△▽▼◎◯○●◉・･]+', ' ', text)
    text = re.sub(r'[~～]{2,}', '…', text)
    text = re.sub(r'[—–-]{3,}', '——', text)
    text = re.sub(r'[!！]{3,}', '!!', text)
    text = re.sub(r'[?？]{3,}', '??', text)
    text = re.sub(r'[…。…]{2,}', '…', text)
    text = re.sub(r'([ぁ-んァ-ヶー])\1{3,}', r'\1\1', text)
    text = re.sub(r'([a-zA-Z])\1{3,}', r'\1\1', text)
    text = re.sub(r'([\u4e00-\u9fff])\1{3,}', r'\1\1', text)
    text = re.sub(r'(?<=[\u3040-\u30ff\u4e00-\u9fff])\s+(?=[、。！？…」』）])', '', text)
    text = re.sub(r'(?<=[（「『])\s+', '', text)
    text = re.sub(r'\s+(?=[」』）])', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    cleaned_lines = []
    for line in text.splitlines():
        line = re.sub(r'\s+', ' ', line).strip()
        if not line:
            if cleaned_lines and cleaned_lines[-1] != '':
                cleaned_lines.append('')
            continue
        symbol_ratio = sum(1 for ch in line if not ch.isalnum() and ch not in ' 一-龥ぁ-んァ-ヶ。，、！？…：；（）()「」『』【】〈〉《》“”"\'')
        if len(line) <= 8 and symbol_ratio >= max(3, len(line) - 1):
            continue
        cleaned_lines.append(line)

    text = '\n'.join(cleaned_lines).strip()
    text = collapse_duplicate_lines(text)
    text = dedupe_repeated_prefix_progression(text)
    text = dedupe_repeated_sentences(text)
    text = re.sub(r'\n\s*\n+', '\n\n', text)
    return text


def split_text_blocks(text: str):
    parts = re.split(r'\n\s*\n+', text.strip())
    out = []
    idx = 1
    for p in parts:
        p = heavy_cleanup_text(p)
        if not p:
            continue
        for piece in split_long_block(p):
            piece = heavy_cleanup_text(piece)
            if not piece:
                continue
            out.append({'id': idx, 'start': None, 'end': None, 'text': piece})
            idx += 1
    return out


def normalize_text(text: str):
    segs = split_text_blocks(heavy_cleanup_text(text))
    return {
        'source_type': 'text',
        'language': 'auto',
        'has_timestamps': False,
        'normalization_plan': 'b',
        'segments': segs,
    }


def normalize_transcript(obj):
    segments = obj.get('segments', [])
    out = []
    idx = 1
    for s in segments:
        text = heavy_cleanup_text((s.get('text') or '').strip())
        if not text:
            continue
        out.append({
            'id': s.get('id', idx),
            'start': s.get('start'),
            'end': s.get('end'),
            'text': text,
        })
        idx += 1
    return {
        'source_type': 'transcript',
        'language': obj.get('language') or obj.get('source_lang') or 'auto',
        'has_timestamps': any(x.get('start') is not None and x.get('end') is not None for x in out),
        'normalization_plan': 'b',
        'segments': out,
    }


def normalize_web(source: str):
    if re.match(r'^https?://', source):
        req = urllib.request.Request(source, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=30) as r:
            html = r.read().decode('utf-8', errors='replace')
    else:
        html = Path(source).read_text(encoding='utf-8', errors='replace')
    html = re.sub(r'(?is)<script.*?>.*?</script>', ' ', html)
    html = re.sub(r'(?is)<style.*?>.*?</style>', ' ', html)
    text = re.sub(r'(?s)<[^>]+>', '\n', html)
    text = unescape(text)
    text = heavy_cleanup_text(text)
    return {
        'source_type': 'web',
        'language': 'auto',
        'has_timestamps': False,
        'normalization_plan': 'b',
        'segments': split_text_blocks(text),
    }


def normalize_pdf(path: str):
    p = Path(path).expanduser().resolve()
    try:
        from pypdf import PdfReader
    except Exception:
        raise SystemExit('PDF support requires pypdf in media-pipeline venv')
    reader = PdfReader(str(p))
    texts = []
    for page in reader.pages:
        t = heavy_cleanup_text((page.extract_text() or '').strip())
        if t:
            texts.append(t)
    return {
        'source_type': 'pdf',
        'language': 'auto',
        'has_timestamps': False,
        'normalization_plan': 'b',
        'segments': split_text_blocks('\n\n'.join(texts)),
    }


def normalize_docx(path: str):
    p = Path(path).expanduser().resolve()
    try:
        from docx import Document
    except Exception:
        raise SystemExit('DOCX support requires python-docx in media-pipeline venv')

    doc = Document(str(p))
    texts = []

    for para in doc.paragraphs:
        t = heavy_cleanup_text((para.text or '').strip())
        if t:
            texts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                t = '\n'.join(heavy_cleanup_text(p.text.strip()) for p in cell.paragraphs if heavy_cleanup_text((p.text or '').strip()))
                t = heavy_cleanup_text(t)
                if t:
                    cells.append(t)
            if cells:
                texts.append(' | '.join(cells))

    return {
        'source_type': 'docx',
        'language': 'auto',
        'has_timestamps': False,
        'normalization_plan': 'b',
        'segments': split_text_blocks('\n\n'.join(texts)),
    }


def normalize_generic_json(obj):
    if isinstance(obj, dict) and isinstance(obj.get('segments'), list):
        return normalize_transcript(obj)

    candidates = None
    if isinstance(obj, dict):
        for key in ['blocks', 'items', 'lines', 'paragraphs', 'texts']:
            if isinstance(obj.get(key), list):
                candidates = obj[key]
                break
    elif isinstance(obj, list):
        candidates = obj

    if candidates is None:
        return normalize_text(json.dumps(obj, ensure_ascii=False, indent=2))

    out = []
    idx = 1
    for item in candidates:
        if isinstance(item, str):
            text = heavy_cleanup_text(item.strip())
            start = end = None
        elif isinstance(item, dict):
            text = heavy_cleanup_text((item.get('text') or item.get('content') or item.get('value') or '').strip())
            start = item.get('start')
            end = item.get('end')
        else:
            text = heavy_cleanup_text(str(item).strip())
            start = end = None
        if not text:
            continue
        out.append({'id': idx, 'start': start, 'end': end, 'text': text})
        idx += 1

    return {
        'source_type': 'json',
        'language': 'auto',
        'has_timestamps': any(x.get('start') is not None and x.get('end') is not None for x in out),
        'normalization_plan': 'b',
        'segments': out,
    }


def normalize_ocr(obj):
    result = normalize_generic_json(obj)
    result['source_type'] = 'ocr'
    return result


def main():
    p = argparse.ArgumentParser(description='Normalize various inputs into segment JSON (plan B)')
    p.add_argument('--input', required=True, help='Input file path or URL')
    p.add_argument('--type', choices=['auto', 'text', 'transcript', 'json', 'web', 'pdf', 'docx', 'ocr'], default='auto')
    p.add_argument('--output', default=None, help='Output JSON path')
    p.add_argument('--json-stdout', action='store_true', help='Print normalized JSON to stdout')
    args = p.parse_args()

    is_url = bool(re.match(r'^https?://', args.input))
    path = None if is_url else Path(args.input).expanduser().resolve()
    if not is_url and not path.exists():
        raise SystemExit(f'Input not found: {path}')

    if args.type == 'auto':
        if is_url:
            input_type = 'web'
        elif path.suffix.lower() == '.json':
            input_type = 'json'
        elif path.suffix.lower() == '.pdf':
            input_type = 'pdf'
        elif path.suffix.lower() == '.docx':
            input_type = 'docx'
        elif path.suffix.lower() in ('.html', '.htm'):
            input_type = 'web'
        else:
            input_type = 'text'
    else:
        input_type = args.type

    if input_type == 'text':
        normalized = normalize_text(path.read_text(encoding='utf-8', errors='replace'))
    elif input_type == 'transcript':
        obj = json.loads(path.read_text(encoding='utf-8'))
        normalized = normalize_transcript(obj)
    elif input_type == 'json':
        obj = json.loads(path.read_text(encoding='utf-8'))
        normalized = normalize_generic_json(obj)
    elif input_type == 'ocr':
        obj = json.loads(path.read_text(encoding='utf-8'))
        normalized = normalize_ocr(obj)
    elif input_type == 'pdf':
        normalized = normalize_pdf(str(path))
    elif input_type == 'docx':
        normalized = normalize_docx(str(path))
    elif input_type == 'web':
        normalized = normalize_web(args.input if is_url else str(path))
    else:
        raise SystemExit(f'Unsupported type: {input_type}')

    if args.output:
        out = Path(args.output).expanduser().resolve()
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(json.dumps(normalized, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')

    if args.json_stdout or not args.output:
        print(json.dumps(normalized, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
